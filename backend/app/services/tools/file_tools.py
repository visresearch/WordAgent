"""Project file tools: list/read/edit under wence_data/project only."""

from __future__ import annotations

import re
from pathlib import Path

from langchain_core.tools import tool
from langgraph.config import get_stream_writer

from app.core.config import get_wence_project_dir

_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".tiff", ".tif"}
_EDITABLE_TEXT_EXTENSIONS = {
    ".md",
    ".txt",
    ".json",
    ".yaml",
    ".yml",
    ".csv",
    ".log",
    ".html",
    ".js",
    ".tx",
    ".py",
    ".mdx",
}
_MAX_LIST_ENTRIES = 500
_MAX_READ_CHARS = 80_000
_LINE_RANGE_RE = re.compile(r"^(.*?):(\d+)(?:-(\d+))?$")
_OCR_ENGINE = None
_OCR_ENGINE_LOAD_ERROR = None


def _project_root() -> Path:
    root = get_wence_project_dir().resolve()
    root.mkdir(parents=True, exist_ok=True)
    return root


def _ensure_inside_project(path: Path) -> None:
    root = _project_root()
    try:
        path.relative_to(root)
    except Exception as exc:
        raise ValueError(f"Path is outside project root: {path}") from exc


def _resolve_project_path(raw_path: str, must_exist: bool = False) -> Path:
    if not raw_path or not raw_path.strip():
        raise ValueError("Path cannot be empty.")

    root = _project_root()
    candidate = Path(raw_path.strip())
    resolved = candidate.resolve(strict=False) if candidate.is_absolute() else (root / candidate).resolve(strict=False)
    _ensure_inside_project(resolved)

    if must_exist and not resolved.exists():
        raise FileNotFoundError(f"Path does not exist: {resolved}")

    return resolved


def _safe_relative(path: Path) -> str:
    root = _project_root()
    return path.relative_to(root).as_posix() if path != root else "."


def _truncate_text(text: str, limit: int = _MAX_READ_CHARS) -> str:
    if len(text) <= limit:
        return text
    keep_head = int(limit * 0.7)
    keep_tail = limit - keep_head
    omitted = len(text) - limit
    return f"{text[:keep_head]}\n\n...[truncated {omitted} chars]...\n\n{text[-keep_tail:]}"


def _read_text_with_fallback(file_path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "gbk"):
        try:
            return file_path.read_text(encoding=enc)
        except UnicodeDecodeError:
            continue
    return file_path.read_text(encoding="utf-8", errors="replace")


def _parse_line_range(path: str, start_line: int | None, end_line: int | None) -> tuple[str, int | None, int | None]:
    if start_line is not None or end_line is not None:
        return path, start_line, end_line

    m = _LINE_RANGE_RE.match(path.strip())
    if not m:
        return path, None, None

    candidate_path = m.group(1).strip()
    s = int(m.group(2))
    e = int(m.group(3)) if m.group(3) else s
    return candidate_path, s, e


def _format_with_line_numbers(text: str, start_line: int | None, end_line: int | None) -> str:
    lines = text.splitlines()
    if not lines:
        return "File is empty."

    if start_line is None:
        start_idx = 0
        end_idx = len(lines)
    else:
        s = max(1, start_line)
        e = max(s, end_line) if end_line is not None else s
        start_idx = min(s - 1, len(lines))
        end_idx = min(e, len(lines))
        if start_idx >= len(lines):
            return f"Requested range starts beyond EOF (total lines: {len(lines)})."

    selected = lines[start_idx:end_idx]
    numbered = [f"{idx}|{line}" for idx, line in enumerate(selected, start=start_idx + 1)]
    return "\n".join(numbered)


def _get_python_ocr_engine():
    """Lazy-load pure Python OCR engine (RapidOCR)."""
    global _OCR_ENGINE, _OCR_ENGINE_LOAD_ERROR
    if _OCR_ENGINE is not None:
        return _OCR_ENGINE
    if _OCR_ENGINE_LOAD_ERROR is not None:
        return None

    try:
        import importlib

        rapidocr_module = importlib.import_module("rapidocr_onnxruntime")
        rapidocr_class = getattr(rapidocr_module, "RapidOCR", None)
        if rapidocr_class is None:
            raise RuntimeError("rapidocr_onnxruntime.RapidOCR not found")
        _OCR_ENGINE = rapidocr_class()
        return _OCR_ENGINE
    except Exception as exc:
        _OCR_ENGINE_LOAD_ERROR = str(exc)
        return None


def _ocr_image_text(file_path: Path) -> str:
    engine = _get_python_ocr_engine()
    if engine is None:
        reason = _OCR_ENGINE_LOAD_ERROR or "rapidocr_onnxruntime is unavailable"
        return (
            f"OCR unavailable: pure Python engine not ready. Please install `rapidocr-onnxruntime`. Details: {reason}"
        )

    try:
        result, _elapsed = engine(str(file_path))
    except Exception as exc:
        return f"OCR failed: {exc}"

    if not result:
        return "OCR completed but no readable text was detected."

    texts: list[str] = []
    for item in result:
        if not isinstance(item, (list, tuple)) or len(item) < 2:
            continue
        text_part = item[1]
        if isinstance(text_part, str):
            cleaned = text_part.strip()
            if cleaned:
                texts.append(cleaned)

    if not texts:
        return "OCR completed but no readable text was detected."
    return _truncate_text("\n".join(texts))


def _get_image_metadata(file_path: Path) -> str:
    """Return image size/format metadata if available."""
    size_bytes = file_path.stat().st_size if file_path.exists() else 0
    try:
        from PIL import Image

        with Image.open(file_path) as img:
            width, height = img.size
            fmt = img.format or "unknown"
            mode = img.mode or "unknown"
        return f"width={width}, height={height}, format={fmt}, mode={mode}, file_size={size_bytes} bytes"
    except Exception:
        return f"file_size={size_bytes} bytes (image metadata unavailable)"


def _extract_pdf_text(file_path: Path) -> str:
    try:
        import importlib

        fitz = importlib.import_module("fitz")  # pymupdf

        doc = fitz.open(str(file_path))
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        text = "\n".join(pages).strip()
        return text if text else "[PDF extracted no readable text.]"
    except Exception as exc:
        return f"[PDF text extraction failed: {exc}]"


def _extract_docx_text(file_path: Path) -> str:
    try:
        import docx

        doc = docx.Document(str(file_path))
        lines = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
        return "\n".join(lines).strip()
    except Exception as exc:
        return f"[DOCX text extraction failed: {exc}]"


def build_list_file(description: str):
    @tool(description=description)
    def list_file(path: str = ".", recursive: bool = False, include_hidden: bool = False) -> str:
        """List files under wence_data/project. Path cannot escape project root."""
        writer = get_stream_writer()
        if writer:
            writer(
                {
                    "type": "status",
                    "content": f"📁 查看文件列表: {path}",
                }
            )
        target = _resolve_project_path(path, must_exist=True)

        if target.is_file():
            stat = target.stat()
            rel = _safe_relative(target)
            return f"[FILE] {rel} ({stat.st_size} bytes)"

        if not target.is_dir():
            return f"Not a file or directory: {_safe_relative(target)}"

        items = target.rglob("*") if recursive else target.iterdir()
        entries: list[str] = []
        for item in sorted(items, key=lambda p: p.as_posix().lower()):
            rel = _safe_relative(item)
            name = item.name
            if not include_hidden and name.startswith("."):
                continue
            if item.is_dir():
                entries.append(f"[DIR]  {rel}/")
            elif item.is_file():
                size = item.stat().st_size
                entries.append(f"[FILE] {rel} ({size} bytes)")
            if len(entries) >= _MAX_LIST_ENTRIES:
                entries.append(f"... truncated at {_MAX_LIST_ENTRIES} entries ...")
                break

        if not entries:
            return f"No entries found in {_safe_relative(target)}."
        return "\n".join(entries)

    return list_file


def build_read_file(description: str):
    @tool(description=description)
    def read_file(path: str, start_line: int | None = None, end_line: int | None = None) -> str:
        """
        Read file content from wence_data/project. Supports optional line range.
        For images, returns OCR text.
        """
        writer = get_stream_writer()
        if writer:
            range_desc = (
                f"start_line={start_line}, end_line={end_line}"
                if (start_line is not None or end_line is not None)
                else "line_range=auto"
            )
            writer(
                {
                    "type": "status",
                    "content": f"✅ 读取文件: {path}",
                }
            )
        raw_path, range_start, range_end = _parse_line_range(path, start_line, end_line)
        file_path = _resolve_project_path(raw_path, must_exist=True)

        if not file_path.is_file():
            return f"Not a file: {_safe_relative(file_path)}"

        suffix = file_path.suffix.lower()
        if suffix in _IMAGE_EXTENSIONS:
            text = _ocr_image_text(file_path)
            image_meta = _get_image_metadata(file_path)
            return f"[IMAGE] {_safe_relative(file_path)}\n[Image Meta] {image_meta}\n[OCR]\n{text}"

        if suffix == ".pdf":
            text = _extract_pdf_text(file_path)
            return f"[PDF] {_safe_relative(file_path)}\n{_truncate_text(_format_with_line_numbers(text, range_start, range_end))}"

        if suffix == ".docx":
            text = _extract_docx_text(file_path)
            return f"[DOCX] {_safe_relative(file_path)}\n{_truncate_text(_format_with_line_numbers(text, range_start, range_end))}"

        text = _read_text_with_fallback(file_path)
        formatted = _format_with_line_numbers(text, range_start, range_end)
        return _truncate_text(formatted)

    return read_file


def build_edit_file(description: str):
    @tool(description=description)
    def edit_file(path: str, content: str, append: bool = False) -> str:
        """
        Create or edit text files under wence_data/project.
        Missing files are created automatically.
        """
        writer = get_stream_writer()
        if writer:
            writer(
                {
                    "type": "status",
                    "content": f"✅ 编辑文件: {path}",
                }
            )
        file_path = _resolve_project_path(path, must_exist=False)

        existed_before = file_path.exists()

        if file_path.exists() and file_path.is_dir():
            return f"Cannot edit directory: {_safe_relative(file_path)}"

        suffix = file_path.suffix.lower()
        if suffix and suffix not in _EDITABLE_TEXT_EXTENSIONS:
            return f"Unsupported file type for edit: {suffix}. Allowed: {sorted(_EDITABLE_TEXT_EXTENSIONS)}"

        file_path.parent.mkdir(parents=True, exist_ok=True)
        mode = "a" if append else "w"
        with open(file_path, mode, encoding="utf-8") as f:
            f.write(content)

        action = "appended" if append else ("created" if not existed_before else "written")
        size = file_path.stat().st_size
        return f"File {action}: {_safe_relative(file_path)} ({size} bytes)"

    return edit_file


__all__ = ["build_list_file", "build_read_file", "build_edit_file"]
