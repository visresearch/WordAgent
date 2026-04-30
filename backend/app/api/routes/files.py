"""
文件上传路由

提供附件上传、查询、清理接口。
上传的文件保存在 wence_data/uploads/ 目录下，以 UUID 命名防止冲突。
"""

import base64
import mimetypes
import uuid
from pathlib import Path

from fastapi import APIRouter, UploadFile, File
from pydantic import BaseModel

from app.core.config import get_wence_data_dir

router = APIRouter()

# 上传目录
UPLOAD_DIR = get_wence_data_dir() / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# 允许的文件扩展名
ALLOWED_EXTENSIONS = {".png", ".jpg", ".jpeg", ".pdf", ".docx", ".txt", ".md"}

# 图片 MIME 类型
IMAGE_MIME_TYPES = {"image/png", "image/jpeg", "image/jpg"}

# 单文件最大 50MB
MAX_FILE_SIZE = 50 * 1024 * 1024

# 注入到 LLM 上下文的文本最大字符数（约 25K tokens），从配置读取
from app.core.config import settings

MAX_TEXT_CHARS = settings.MAX_TEXT_CHARS


# ============== 响应模型 ==============


class UploadedFileInfo(BaseModel):
    """单个上传文件的信息"""

    file_id: str
    filename: str
    size: int
    content_type: str
    is_image: bool


class UploadResponse(BaseModel):
    """上传响应"""

    success: bool
    files: list[UploadedFileInfo] = []
    error: str | None = None


# ============== 工具函数 ==============


def _is_allowed_extension(filename: str) -> bool:
    """检查文件扩展名是否在允许列表中"""
    suffix = Path(filename).suffix.lower()
    return suffix in ALLOWED_EXTENSIONS


def _get_content_type(filename: str) -> str:
    """获取文件 MIME 类型"""
    mime, _ = mimetypes.guess_type(filename)
    return mime or "application/octet-stream"


def get_file_path(file_id: str) -> Path | None:
    """根据 file_id 获取文件路径（安全校验）"""
    # file_id 格式: uuid_originalname
    file_path = UPLOAD_DIR / file_id
    if file_path.exists() and file_path.is_file() and file_path.resolve().parent == UPLOAD_DIR.resolve():
        return file_path
    return None


def read_file_as_base64(file_id: str) -> str | None:
    """读取文件并返回 base64 编码"""
    file_path = get_file_path(file_id)
    if not file_path:
        return None
    with open(file_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")


def _truncate_text(text: str, max_chars: int = MAX_TEXT_CHARS) -> str:
    """截断过长的文本，保留头尾"""
    if len(text) <= max_chars:
        return text
    # 头部保留 80%，尾部保留 20%
    head_size = int(max_chars * 0.8)
    tail_size = max_chars - head_size
    return (
        text[:head_size] + f"\n\n... [内容过长，已省略中间 {len(text) - max_chars} 个字符] ...\n\n" + text[-tail_size:]
    )


def read_text_file(file_id: str, original_filename: str) -> str | None:
    """读取文本类文件内容（自动截断过长内容）"""
    file_path = get_file_path(file_id)
    if not file_path:
        return None

    suffix = Path(original_filename).suffix.lower()
    text = None

    if suffix in (".txt", ".md"):
        try:
            text = file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = file_path.read_text(encoding="gbk", errors="replace")

    elif suffix == ".pdf":
        text = _extract_pdf_text(file_path)

    elif suffix == ".docx":
        text = _extract_docx_text(file_path)

    if text:
        return _truncate_text(text)
    return None


def _extract_pdf_text(file_path: Path) -> str:
    """从 PDF 提取文本（使用 pymupdf）"""
    try:
        import fitz  # pymupdf

        doc = fitz.open(str(file_path))
        pages = []
        for page in doc:
            pages.append(page.get_text())
        doc.close()
        text = "\n".join(pages).strip()
        return text if text else "[PDF 文件内容为空（可能是扫描件/纯图片 PDF）]"
    except Exception as e:
        return f"[PDF 文件，文本提取失败: {e}]"


def _extract_docx_text(file_path: Path) -> str:
    """从 DOCX 提取纯文本（使用 python-docx）"""
    try:
        import docx

        doc = docx.Document(str(file_path))
        lines = [p.text for p in doc.paragraphs]

        # 同时提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))

        return "\n".join(lines)
    except Exception as e:
        return f"[DOCX 文件，文本提取失败: {e}]"


# ============== 路由 ==============


@router.post("/upload", response_model=UploadResponse)
async def upload_files(files: list[UploadFile] = File(...)):
    """
    上传一个或多个文件

    支持格式: png, jpg, jpeg, pdf, docx, txt, md
    单文件最大 50MB
    返回每个文件的 file_id，供后续 chat 请求引用
    """
    results: list[UploadedFileInfo] = []

    for upload_file in files:
        filename = upload_file.filename or "unknown"

        # 校验扩展名
        if not _is_allowed_extension(filename):
            continue

        # 读取文件内容
        content = await upload_file.read()

        # 校验大小
        if len(content) > MAX_FILE_SIZE:
            continue

        # 生成唯一文件名
        file_id = f"{uuid.uuid4().hex}_{filename}"
        save_path = UPLOAD_DIR / file_id
        save_path.write_bytes(content)

        content_type = _get_content_type(filename)
        is_image = content_type in IMAGE_MIME_TYPES

        results.append(
            UploadedFileInfo(
                file_id=file_id,
                filename=filename,
                size=len(content),
                content_type=content_type,
                is_image=is_image,
            )
        )

    return UploadResponse(success=True, files=results)


@router.delete("/upload/{file_id}")
async def delete_uploaded_file(file_id: str):
    """删除已上传的文件"""
    file_path = get_file_path(file_id)
    if file_path:
        file_path.unlink(missing_ok=True)
        return {"success": True}
    return {"success": False, "error": "文件不存在"}
